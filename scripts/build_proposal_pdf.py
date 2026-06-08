from pathlib import Path

from docx import Document
from docx.shared import Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "competition"
PDF_PATH = OUT_DIR / "MoonDepSolve项目申报书.pdf"
DOCX_PATH = OUT_DIR / "MoonDepSolve项目申报书.docx"

TITLE = "MoonDepSolve 项目申报书"
PROJECT_NAME = "MoonDepSolve：MoonBit 包生态语义版本与依赖求解工具"
GITLINK_URL = "https://gitlink.org.cn/python123/moondepsolve"
GITHUB_URL = "https://github.com/python123-ops/moondepsolve"

SECTIONS = [
    (
        "一、项目简介",
        "MoonDepSolve 面向 MoonBit 包生态中的版本约束和依赖选择问题，提供语义版本解析、版本范围匹配、依赖图求解、冲突诊断和依赖锁定结果输出。项目使用 MoonBit 编写，不依赖外部服务，适合在包管理、构建规划、依赖检查和自动化发布流程中作为基础组件使用。",
    ),
    (
        "二、项目方向与适用场景",
        "依赖求解是包管理和构建系统中的基础能力。随着 MoonBit 包数量增加，工具链需要能够判断版本兼容性、选择可用版本组合、解释依赖冲突，并形成稳定的解析结果。MoonDepSolve 可服务于依赖安装、构建规划、兼容性检查、依赖升级建议和教学示例等场景。",
    ),
    (
        "三、计划实现的核心功能",
        "项目目前实现语义版本解析、prerelease 排序、exact/caret/tilde/comparator/wildcard 五类约束解析、版本匹配、传递依赖求解、最高兼容版本选择、冲突路径诊断和依赖锁定结果输出。命令行示例会展示一个内置包索引的求解过程，测试覆盖成功解析和失败诊断场景。",
    ),
    (
        "四、实现说明",
        "项目代码围绕 MoonBit 版本约束场景自行设计数据结构、解析器、求解流程和错误诊断。目前版本控制在较小规模内，不引入外部依赖，也不绑定具体平台接口，便于评审直接克隆、运行和检查。",
    ),
    (
        "五、技术路线",
        "项目采用“解析层、约束层、求解层、展示层”的结构。解析层负责版本和约束字符串；约束层将范围统一表示为 comparator 数组；求解层从根依赖开始遍历包索引，选择最高兼容版本并展开传递依赖；诊断层在包缺失、无匹配版本或已选版本冲突时输出可读错误；展示层通过 CLI 和依赖锁定结果文本展示结果。",
    ),
    (
        "六、预期成果",
        "项目预期交付一个可运行的 MoonBit 依赖求解基础库，一组覆盖版本解析、约束匹配、传递依赖和冲突诊断的测试，一个可通过 moon run cmd/main 执行的 CLI 示例，以及 README、设计文档、验收清单和本申报书。仓库将同步到 GitLink 与 GitHub，便于评审克隆和复现。",
    ),
    (
        "七、后续计划",
        "后续计划接入真实包索引读取、结构化锁文件读写、最小变更升级建议、依赖图输出、冲突图解释、构建计划生成和 MoonBit 包发布流程。相比单一文本处理工具，依赖求解方向更贴近包生态和构建工具链中的基础需求。",
    ),
]


def register_font() -> str:
    candidates = [
        Path("C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf"),
        Path("C:/Windows/Fonts/simsun.ttc"),
    ]
    for font in candidates:
        if font.exists():
            pdfmetrics.registerFont(TTFont("MoonDepSolveCN", str(font)))
            return "MoonDepSolveCN"
    return "Helvetica"


def build_pdf() -> None:
    font = register_font()
    styles = getSampleStyleSheet()
    title = ParagraphStyle(
        "TitleCN",
        parent=styles["Title"],
        fontName=font,
        fontSize=22,
        leading=30,
        alignment=TA_CENTER,
        spaceAfter=20,
    )
    heading = ParagraphStyle(
        "HeadingCN",
        parent=styles["Heading2"],
        fontName=font,
        fontSize=13,
        leading=18,
        textColor=colors.HexColor("#1f3b73"),
        spaceBefore=10,
        spaceAfter=6,
    )
    body = ParagraphStyle(
        "BodyCN",
        parent=styles["BodyText"],
        fontName=font,
        fontSize=10.5,
        leading=17,
        firstLineIndent=21,
        spaceAfter=5,
    )
    meta = ParagraphStyle(
        "MetaCN",
        parent=styles["BodyText"],
        fontName=font,
        fontSize=10.5,
        leading=16,
    )

    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=A4,
        rightMargin=2 * cm,
        leftMargin=2 * cm,
        topMargin=1.8 * cm,
        bottomMargin=1.8 * cm,
        title=TITLE,
    )
    story = [Paragraph(TITLE, title)]
    table = Table(
        [
            [Paragraph("项目名称", meta), Paragraph(PROJECT_NAME, meta)],
            [Paragraph("参赛方向", meta), Paragraph("MoonBit 国产基础软件开源生态项目", meta)],
            [Paragraph("开源许可证", meta), Paragraph("Apache-2.0", meta)],
            [Paragraph("GitLink 仓库", meta), Paragraph(GITLINK_URL, meta)],
            [Paragraph("GitHub 仓库", meta), Paragraph(GITHUB_URL, meta)],
        ],
        colWidths=[3.2 * cm, 12 * cm],
    )
    table.setStyle(
        TableStyle(
            [
                ("FONTNAME", (0, 0), (-1, -1), font),
                ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#edf3ff")),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#9aa9c7")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]
        )
    )
    story.extend([table, Spacer(1, 0.4 * cm)])
    for section_title, text in SECTIONS:
        story.append(Paragraph(section_title, heading))
        story.append(Paragraph(text, body))
    doc.build(story)


def build_docx() -> None:
    doc = Document()
    doc.styles["Normal"].font.name = "Microsoft YaHei"
    doc.styles["Normal"].font.size = Pt(10.5)
    title = doc.add_heading(TITLE, level=0)
    title.alignment = 1
    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    rows = [
        ("项目名称", PROJECT_NAME),
        ("参赛方向", "MoonBit 国产基础软件开源生态项目"),
        ("开源许可证", "Apache-2.0"),
        ("GitLink 仓库", GITLINK_URL),
        ("GitHub 仓库", GITHUB_URL),
    ]
    for row, (key, value) in zip(table.rows, rows):
        row.cells[0].text = key
        row.cells[1].text = value
    for section_title, text in SECTIONS:
        doc.add_heading(section_title, level=1)
        doc.add_paragraph(text)
    doc.save(DOCX_PATH)


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    build_pdf()
    build_docx()
    print(PDF_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
